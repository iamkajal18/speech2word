from pydub import AudioSegment
import speech_recognition as sr
from docx import Document

audio = AudioSegment.from_mp3("my_audio.mp3")
audio.export("c_audio.wav", format="wav")

r = sr.Recognizer()
with sr.AudioFile("c_audio.wav") as source:
    data = r.record(source)

try:
    text = r.recognize_google(data)
    print("Text:", text)

    doc = Document()
    doc.add_heading("Transcription", 0)
    doc.add_paragraph(text)
    doc.save("output.docx")
    print("Saved as output.docx")

except sr.UnknownValueError:
    print("Audio samajh nahi aayi.")
except sr.RequestError as e:
    print("Error:", e)
