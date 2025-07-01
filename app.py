import os
from flask import Flask, render_template, request, send_file
from pydub import AudioSegment
import speech_recognition as sr
from docx import Document

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files['file']
        if file.filename == '':
            return " No file selected."

        
        mp3_path = os.path.join(UPLOAD_FOLDER, 'my_audio.mp3')
        wav_path = os.path.join(UPLOAD_FOLDER, 'c_audio.wav')
        docx_path = 'transcription.docx'

      
        file.save(mp3_path)

        try:
           
            audio = AudioSegment.from_mp3(mp3_path)
            audio = audio.set_channels(1).set_frame_rate(16000)  
            audio.export(wav_path, format='wav')

            
            recognizer = sr.Recognizer()
            with sr.AudioFile(wav_path) as source:
                recognizer.adjust_for_ambient_noise(source)  
                audio_data = recognizer.record(source)

            try:
                text = recognizer.recognize_google(audio_data)
            except sr.UnknownValueError:
                text = "not clear ."
            except sr.RequestError as e:
                text = f" no response: {e}"

         
            doc = Document()
            doc.add_heading("Audio Transcription", 0)
            doc.add_paragraph(text)
            doc.save(docx_path)

            return send_file(docx_path, as_attachment=True)

        except Exception as e:
            return f" Error: {str(e)}"

    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)
